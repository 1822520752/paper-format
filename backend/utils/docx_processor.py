import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from typing import Dict, List, Any, Tuple
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxProcessor:
    """Word文档处理核心类"""
    
    FONT_SIZE_MAP = {
        "初号": 42, "小初": 36, "一号": 26, "小一": 24,
        "二号": 22, "小二": 18, "三号": 16, "小三": 15,
        "四号": 14, "小四": 12, "五号": 10.5, "小五": 9
    }
    
    HEADING_PATTERNS = {
        1: [
            r'^第[一二三四五六七八九十百]+章',
            r'^第\d+章',
            r'^摘\s*要$',
            r'^Abstract$',
            r'^目\s*录$',
            r'^参考文献$',
            r'^致\s*谢$',
            r'^附\s*录$'
        ],
        2: [r'^\d+\.\d+\s+', r'^\d+\.\d+$'],
        3: [r'^\d+\.\d+\.\d+\s+', r'^\d+\.\d+\.\d+$']
    }
    
    def __init__(self, file_path: str):
        """初始化文档处理器"""
        try:
            self.doc = Document(file_path)
            self.file_path = file_path
            logger.info(f"成功加载文档: {file_path}")
        except Exception as e:
            logger.error(f"加载文档失败: {str(e)}")
            raise ValueError(f"无法打开文档: {str(e)}")
    
    def check_format(self, config: Dict[str, Any]) -> Dict[str, Any]:
        """检查文档格式并生成报告"""
        try:
            report = {
                "total_items": 0,
                "passed_items": 0,
                "failed_items": 0,
                "pass_rate": 0,
                "items": []
            }
            
            report["items"].extend(self._check_page_settings(config))
            report["items"].extend(self._check_cover())
            report["items"].extend(self._check_abstract(config))
            report["items"].extend(self._check_toc())
            report["items"].extend(self._check_headings(config))
            report["items"].extend(self._check_body(config))
            report["items"].extend(self._check_figures(config))
            report["items"].extend(self._check_header_footer(config))
            report["items"].extend(self._check_references(config))
            
            report["total_items"] = len(report["items"])
            report["passed_items"] = sum(1 for item in report["items"] if item["passed"])
            report["failed_items"] = report["total_items"] - report["passed_items"]
            report["pass_rate"] = round((report["passed_items"] / report["total_items"] * 100) if report["total_items"] > 0 else 0, 1)
            
            return report
        except Exception as e:
            logger.error(f"格式检查失败: {str(e)}")
            raise
    
    def _check_page_settings(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查页面设置"""
        items = []
        
        if not self.doc.sections or len(self.doc.sections) == 0:
            items.append({
                "category": "页面设置",
                "name": "文档节",
                "passed": False,
                "current": "未找到",
                "expected": "存在",
                "suggestion": "文档缺少节设置"
            })
            return items
        
        section = self.doc.sections[0]
        page_config = config.get("page_settings", {})
        
        # 检查上页边距
        if section.top_margin is not None:
            top_margin = section.top_margin.cm
            expected_top = page_config.get("top_margin", 2.5)
            items.append({
                "category": "页面设置",
                "name": "上页边距",
                "passed": abs(top_margin - expected_top) < 0.2,
                "current": f"{top_margin:.1f}cm",
                "expected": f"{expected_top}cm",
                "suggestion": f"将上页边距调整为{expected_top}cm"
            })
        else:
            items.append({
                "category": "页面设置",
                "name": "上页边距",
                "passed": False,
                "current": "未设置",
                "expected": f"{page_config.get('top_margin', 2.5)}cm",
                "suggestion": "设置上页边距"
            })
        
        # 检查下页边距
        if section.bottom_margin is not None:
            bottom_margin = section.bottom_margin.cm
            expected_bottom = page_config.get("bottom_margin", 2.5)
            items.append({
                "category": "页面设置",
                "name": "下页边距",
                "passed": abs(bottom_margin - expected_bottom) < 0.2,
                "current": f"{bottom_margin:.1f}cm",
                "expected": f"{expected_bottom}cm",
                "suggestion": f"将下页边距调整为{expected_bottom}cm"
            })
        else:
            items.append({
                "category": "页面设置",
                "name": "下页边距",
                "passed": False,
                "current": "未设置",
                "expected": f"{page_config.get('bottom_margin', 2.5)}cm",
                "suggestion": "设置下页边距"
            })
        
        # 检查左页边距
        if section.left_margin is not None:
            left_margin = section.left_margin.cm
            expected_left = page_config.get("left_margin", 3.0)
            items.append({
                "category": "页面设置",
                "name": "左页边距",
                "passed": abs(left_margin - expected_left) < 0.2,
                "current": f"{left_margin:.1f}cm",
                "expected": f"{expected_left}cm",
                "suggestion": f"将左页边距调整为{expected_left}cm"
            })
        else:
            items.append({
                "category": "页面设置",
                "name": "左页边距",
                "passed": False,
                "current": "未设置",
                "expected": f"{page_config.get('left_margin', 3.0)}cm",
                "suggestion": "设置左页边距"
            })
        
        # 检查右页边距
        if section.right_margin is not None:
            right_margin = section.right_margin.cm
            expected_right = page_config.get("right_margin", 2.5)
            items.append({
                "category": "页面设置",
                "name": "右页边距",
                "passed": abs(right_margin - expected_right) < 0.2,
                "current": f"{right_margin:.1f}cm",
                "expected": f"{expected_right}cm",
                "suggestion": f"将右页边距调整为{expected_right}cm"
            })
        else:
            items.append({
                "category": "页面设置",
                "name": "右页边距",
                "passed": False,
                "current": "未设置",
                "expected": f"{page_config.get('right_margin', 2.5)}cm",
                "suggestion": "设置右页边距"
            })
        
        return items
    
    def _check_cover(self) -> List[Dict[str, Any]]:
        """检查封面页"""
        items = []
        has_cover = len(self.doc.paragraphs) > 0
        items.append({
            "category": "封面页",
            "name": "封面存在性",
            "passed": has_cover,
            "current": "存在" if has_cover else "不存在",
            "expected": "存在",
            "suggestion": "添加封面页" if not has_cover else ""
        })
        return items
    
    def _check_abstract(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查摘要格式"""
        items = []
        abstract_found = False
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^摘\s*要$', text) or text == "Abstract":
                abstract_found = True
                is_chinese = re.match(r'^摘\s*要$', text)
                
                abstract_config = config.get("abstract_title", {}).get("chinese" if is_chinese else "english", {})
                expected_font = abstract_config.get("font_name", "黑体" if is_chinese else "Times New Roman")
                expected_size = abstract_config.get("font_size", 18)
                expected_align = abstract_config.get("alignment", "center")
                
                actual_font = para.runs[0].font.name if para.runs and para.runs[0].font.name else "未知"
                actual_size = para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 0
                actual_align = "center" if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else "left"
                
                items.append({
                    "category": "摘要",
                    "name": f"{'中文' if is_chinese else '英文'}摘要标题字体",
                    "passed": actual_font == expected_font,
                    "current": actual_font,
                    "expected": expected_font,
                    "suggestion": f"将字体调整为{expected_font}"
                })
                
                items.append({
                    "category": "摘要",
                    "name": f"{'中文' if is_chinese else '英文'}摘要标题字号",
                    "passed": abs(actual_size - expected_size) < 1,
                    "current": f"{actual_size}pt",
                    "expected": f"{expected_size}pt",
                    "suggestion": f"将字号调整为{expected_size}pt"
                })
                
                items.append({
                    "category": "摘要",
                    "name": f"{'中文' if is_chinese else '英文'}摘要标题对齐",
                    "passed": actual_align == expected_align,
                    "current": "居中" if actual_align == "center" else "左对齐",
                    "expected": "居中" if expected_align == "center" else "左对齐",
                    "suggestion": f"调整为{'居中' if expected_align == 'center' else '左对齐'}"
                })
        
        if not abstract_found:
            items.append({
                "category": "摘要",
                "name": "摘要存在性",
                "passed": False,
                "current": "未找到",
                "expected": "存在",
                "suggestion": "添加中文和英文摘要"
            })
        
        return items
    
    def _check_toc(self) -> List[Dict[str, Any]]:
        """检查目录"""
        items = []
        toc_found = any(re.match(r'^目\s*录$', para.text.strip()) for para in self.doc.paragraphs)
        
        items.append({
            "category": "目录",
            "name": "目录存在性",
            "passed": toc_found,
            "current": "存在" if toc_found else "不存在",
            "expected": "存在",
            "suggestion": "添加目录页" if not toc_found else ""
        })
        
        return items
    
    def _check_headings(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查标题格式"""
        items = []
        
        for para in self.doc.paragraphs:
            level = self._detect_heading_level(para)
            if level > 0:
                heading_config = config.get(f"heading{level}", {})
                
                if para.runs and len(para.runs) > 0:
                    run_font = para.runs[0].font
                    actual_font = run_font.name if run_font and run_font.name else "未知"
                    actual_size = run_font.size.pt if run_font and run_font.size else 0
                    actual_bold = run_font.bold if run_font else False
                    
                    expected_font = heading_config.get("font_name", "黑体")
                    expected_size = heading_config.get("font_size", 16)
                    expected_bold = heading_config.get("bold", True)
                    expected_align = heading_config.get("alignment", "center")
                    
                    actual_align = "center" if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else "left"
                    
                    items.append({
                        "category": f"{level}级标题",
                        "name": f"字体",
                        "passed": actual_font == expected_font,
                        "current": actual_font,
                        "expected": expected_font,
                        "suggestion": f"将字体调整为{expected_font}"
                    })
                    
                    items.append({
                        "category": f"{level}级标题",
                        "name": f"字号",
                        "passed": abs(actual_size - expected_size) < 1,
                        "current": f"{actual_size}pt",
                        "expected": f"{expected_size}pt",
                        "suggestion": f"将字号调整为{expected_size}pt"
                    })
                    
                    items.append({
                        "category": f"{level}级标题",
                        "name": f"加粗",
                        "passed": actual_bold == expected_bold,
                        "current": "是" if actual_bold else "否",
                        "expected": "是" if expected_bold else "否",
                        "suggestion": f"{'添加' if expected_bold else '取消'}加粗"
                    })
                    
                    items.append({
                        "category": f"{level}级标题",
                        "name": f"对齐方式",
                        "passed": actual_align == expected_align,
                        "current": "居中" if actual_align == "center" else "左对齐",
                        "expected": "居中" if expected_align == "center" else "左对齐",
                        "suggestion": f"调整为{'居中' if expected_align == 'center' else '左对齐'}"
                    })
        
        return items
    
    def _check_body(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查正文格式"""
        items = []
        body_config = config.get("body", {})
        
        checked_paragraphs = 0
        for para in self.doc.paragraphs:
            if self._detect_heading_level(para) == 0 and len(para.text.strip()) > 10:
                checked_paragraphs += 1
                if checked_paragraphs > 5:
                    break
                
                if para.runs and len(para.runs) > 0:
                    run_font = para.runs[0].font
                    actual_font = run_font.name if run_font and run_font.name else "未知"
                    actual_size = run_font.size.pt if run_font and run_font.size else 0
                    
                    expected_font = body_config.get("font_name", "宋体")
                    expected_size = body_config.get("font_size", 12)
                    expected_line_spacing = body_config.get("line_spacing", 1.5)
                    expected_indent = body_config.get("first_line_indent", 2)
                    
                    items.append({
                        "category": "正文",
                        "name": "字体",
                        "passed": actual_font == expected_font,
                        "current": actual_font,
                        "expected": expected_font,
                        "suggestion": f"将字体调整为{expected_font}"
                    })
                    
                    items.append({
                        "category": "正文",
                        "name": "字号",
                        "passed": abs(actual_size - expected_size) < 1,
                        "current": f"{actual_size}pt",
                        "expected": f"{expected_size}pt",
                        "suggestion": f"将字号调整为{expected_size}pt"
                    })
                    
                    actual_indent = para.paragraph_format.first_line_indent
                    indent_chars = actual_indent.cm / 0.37 if actual_indent else 0
                    
                    items.append({
                        "category": "正文",
                        "name": "首行缩进",
                        "passed": abs(indent_chars - expected_indent) < 0.5,
                        "current": f"{indent_chars:.1f}字符",
                        "expected": f"{expected_indent}字符",
                        "suggestion": f"将首行缩进调整为{expected_indent}字符"
                    })
                    
                    break
        
        return items
    
    def _check_figures(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查图表标题"""
        items = []
        figure_config = config.get("figure_caption", {})
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^(图|表)\s*\d+', text):
                if para.runs and len(para.runs) > 0:
                    run_font = para.runs[0].font
                    actual_font = run_font.name if run_font and run_font.name else "未知"
                    actual_size = run_font.size.pt if run_font and run_font.size else 0
                    actual_align = "center" if para.alignment == WD_ALIGN_PARAGRAPH.CENTER else "left"
                    
                    expected_font = figure_config.get("font_name", "宋体")
                    expected_size = figure_config.get("font_size", 10.5)
                    expected_align = figure_config.get("alignment", "center")
                    
                    items.append({
                        "category": "图表标题",
                        "name": "字体",
                        "passed": actual_font == expected_font,
                        "current": actual_font,
                        "expected": expected_font,
                        "suggestion": f"将字体调整为{expected_font}"
                    })
                    
                    items.append({
                        "category": "图表标题",
                        "name": "字号",
                        "passed": abs(actual_size - expected_size) < 1,
                        "current": f"{actual_size}pt",
                        "expected": f"{expected_size}pt",
                        "suggestion": f"将字号调整为{expected_size}pt"
                    })
                    
                    break
        
        return items
    
    def _check_header_footer(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查页眉页脚"""
        items = []
        section = self.doc.sections[0]
        
        has_header = section.header is not None
        items.append({
            "category": "页眉页脚",
            "name": "页眉存在性",
            "passed": has_header,
            "current": "存在" if has_header else "不存在",
            "expected": "存在",
            "suggestion": "添加页眉" if not has_header else ""
        })
        
        has_footer = section.footer is not None
        items.append({
            "category": "页眉页脚",
            "name": "页码存在性",
            "passed": has_footer,
            "current": "存在" if has_footer else "不存在",
            "expected": "存在",
            "suggestion": "添加页码" if not has_footer else ""
        })
        
        return items
    
    def _check_references(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """检查参考文献"""
        items = []
        ref_found = False
        
        for para in self.doc.paragraphs:
            if re.match(r'^参考文献$', para.text.strip()):
                ref_found = True
                ref_config = config.get("reference", {})
                
                if para.runs and len(para.runs) > 0:
                    run_font = para.runs[0].font
                    actual_font = run_font.name if run_font and run_font.name else "未知"
                    actual_size = run_font.size.pt if run_font and run_font.size else 0
                    
                    expected_font = ref_config.get("title_font_name", "黑体")
                    expected_size = ref_config.get("title_font_size", 16)
                    
                    items.append({
                        "category": "参考文献",
                        "name": "标题字体",
                        "passed": actual_font == expected_font,
                        "current": actual_font,
                        "expected": expected_font,
                        "suggestion": f"将字体调整为{expected_font}"
                    })
                    
                    items.append({
                        "category": "参考文献",
                        "name": "标题字号",
                        "passed": abs(actual_size - expected_size) < 1,
                        "current": f"{actual_size}pt",
                        "expected": f"{expected_size}pt",
                        "suggestion": f"将字号调整为{expected_size}pt"
                    })
        
        if not ref_found:
            items.append({
                "category": "参考文献",
                "name": "存在性",
                "passed": False,
                "current": "未找到",
                "expected": "存在",
                "suggestion": "添加参考文献部分"
            })
        
        return items
    
    def _detect_heading_level(self, paragraph) -> int:
        """检测段落的标题级别"""
        if paragraph and paragraph.style and paragraph.style.name:
            if paragraph.style.name.startswith('Heading'):
                try:
                    return int(paragraph.style.name.split()[-1])
                except:
                    pass
        
        text = paragraph.text.strip()
        for level, patterns in self.HEADING_PATTERNS.items():
            for pattern in patterns:
                if re.match(pattern, text):
                    return level
        
        return 0
    
    def format_document(self, config: Dict[str, Any], output_path: str) -> str:
        """一键排版文档"""
        try:
            self._apply_page_settings(config)
            self._apply_heading_formats(config)
            self._apply_body_formats(config)
            self._apply_header_footer(config)
            
            self.doc.save(output_path)
            logger.info(f"文档排版完成: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"文档排版失败: {str(e)}")
            raise
    
    def _apply_page_settings(self, config: Dict[str, Any]):
        """应用页面设置"""
        page_config = config.get("page_settings", {})
        for section in self.doc.sections:
            section.top_margin = Cm(page_config.get("top_margin", 2.5))
            section.bottom_margin = Cm(page_config.get("bottom_margin", 2.5))
            section.left_margin = Cm(page_config.get("left_margin", 3.0))
            section.right_margin = Cm(page_config.get("right_margin", 2.5))
    
    def _apply_heading_formats(self, config: Dict[str, Any]):
        """应用标题格式"""
        for para in self.doc.paragraphs:
            level = self._detect_heading_level(para)
            if level > 0:
                heading_config = config.get(f"heading{level}", {})
                self._set_paragraph_format(para, heading_config)
    
    def _apply_body_formats(self, config: Dict[str, Any]):
        """应用正文格式"""
        body_config = config.get("body", {})
        
        for para in self.doc.paragraphs:
            if self._detect_heading_level(para) == 0 and len(para.text.strip()) > 0:
                self._set_paragraph_format(para, body_config)
                para.paragraph_format.first_line_indent = Cm(body_config.get("first_line_indent", 2) * 0.37)
    
    def _apply_header_footer(self, config: Dict[str, Any]):
        """应用页眉页脚"""
        header_config = config.get("header", {})
        footer_config = config.get("footer", {})
        
        for section in self.doc.sections:
            if header_config.get("content"):
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = header_config.get("content", "")
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for run in header_para.runs:
                    run.font.name = header_config.get("font_name", "宋体")
                    run.font.size = Pt(header_config.get("font_size", 9))
    
    def _set_paragraph_format(self, paragraph, format_config: Dict[str, Any]):
        """设置段落格式"""
        if not paragraph or not paragraph.runs:
            return
        
        for run in paragraph.runs:
            if not run or not run.font:
                continue
            
            try:
                font_name = format_config.get("font_name", "宋体")
                run.font.name = font_name
                if run.element.rPr is not None:
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                
                font_size = format_config.get("font_size", 12)
                run.font.size = Pt(font_size)
                
                if "bold" in format_config:
                    run.font.bold = format_config["bold"]
            except Exception as e:
                logger.warning(f"设置段落格式失败: {str(e)}")
                continue
        
        alignment = format_config.get("alignment", "left")
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        if "line_spacing" in format_config:
            line_spacing = format_config["line_spacing"]
            if format_config.get("line_spacing_type") == "fixed":
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)
            else:
                paragraph.paragraph_format.line_spacing = line_spacing
        
        if "space_before" in format_config:
            paragraph.paragraph_format.space_before = Pt(format_config["space_before"] * 12)
        
        if "space_after" in format_config:
            paragraph.paragraph_format.space_after = Pt(format_config["space_after"] * 12)